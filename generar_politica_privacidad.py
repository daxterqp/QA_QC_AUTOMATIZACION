"""
Genera el documento Word de Política de Privacidad para Flow-QA/QC
"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ─── ESTILOS ───────────────────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

# Heading 1
h1 = doc.styles['Heading 1']
h1.font.name = 'Calibri'
h1.font.size = Pt(18)
h1.font.bold = True
h1.font.color.rgb = RGBColor(0x0D, 0x2B, 0x45)
h1.paragraph_format.space_before = Pt(24)
h1.paragraph_format.space_after = Pt(12)

# Heading 2
h2 = doc.styles['Heading 2']
h2.font.name = 'Calibri'
h2.font.size = Pt(14)
h2.font.bold = True
h2.font.color.rgb = RGBColor(0x0E, 0x74, 0x90)
h2.paragraph_format.space_before = Pt(18)
h2.paragraph_format.space_after = Pt(8)

# Heading 3
h3 = doc.styles['Heading 3']
h3.font.name = 'Calibri'
h3.font.size = Pt(12)
h3.font.bold = True
h3.font.color.rgb = RGBColor(0x1A, 0x4A, 0x7A)
h3.paragraph_format.space_before = Pt(12)
h3.paragraph_format.space_after = Pt(6)

# Márgenes
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)


# ─── HELPERS ───────────────────────────────────────────

NAVY = RGBColor(0x0D, 0x2B, 0x45)
TEAL = RGBColor(0x0E, 0x74, 0x90)
ORANGE = RGBColor(0xF9, 0x73, 0x16)
GRAY = RGBColor(0x33, 0x33, 0x33)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLANK_BG = RGBColor(0xFF, 0xFF, 0xCC)  # amarillo suave para campos a rellenar


def add_blank(paragraph, text="________________________"):
    """Agrega un campo subrayado en amarillo para rellenar."""
    run = paragraph.add_run(text)
    run.font.color.rgb = ORANGE
    run.font.underline = True
    run.font.bold = True
    # Highlight amarillo
    rPr = run._element.get_or_add_rPr()
    highlight = OxmlElement('w:highlight')
    highlight.set(qn('w:val'), 'yellow')
    rPr.append(highlight)
    return run


def add_para(text, bold=False, italic=False, size=None, color=None, align=None, space_after=None):
    """Agrega un párrafo con formato."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.font.bold = True
    if italic:
        run.font.italic = True
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    if align:
        p.alignment = align
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    return p


def add_bullet(text, bold_prefix=None):
    """Agrega un bullet point, opcionalmente con prefijo en negrita."""
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run_b = p.add_run(bold_prefix)
        run_b.font.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def add_bullet_with_blank(prefix, blank_text="________________________"):
    """Bullet con campo a rellenar."""
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(prefix + " ")
    run.font.bold = True
    add_blank(p, blank_text)
    return p


def set_cell_shading(cell, color_hex):
    """Aplica color de fondo a una celda."""
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), color_hex)
    tcPr.append(shading)


def add_table(headers, rows, col_widths=None):
    """Crea una tabla con encabezado navy."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.font.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '0D2B45')

    # Rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            run.font.name = 'Calibri'
            if r_idx % 2 == 1:
                set_cell_shading(cell, 'F1F5F9')

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)

    doc.add_paragraph()  # espacio después
    return table


def add_separator():
    """Línea separadora."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run('─' * 72)
    run.font.size = Pt(6)
    run.font.color.rgb = RGBColor(0xCB, 0xD5, 0xE1)


# ═══════════════════════════════════════════════════════
# CONTENIDO DEL DOCUMENTO
# ═══════════════════════════════════════════════════════

# ─── PORTADA ───────────────────────────────────────────
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('FLOW-QA/QC')
run.font.size = Pt(36)
run.font.bold = True
run.font.color.rgb = NAVY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Sistema Digital de Control de Calidad en Obra')
run.font.size = Pt(14)
run.font.color.rgb = TEAL

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('POLÍTICA DE PRIVACIDAD\nY PROTECCIÓN DE DATOS PERSONALES')
run.font.size = Pt(22)
run.font.bold = True
run.font.color.rgb = NAVY

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('Versión: ').font.bold = True
p.add_run('1.0')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('Fecha de vigencia: ').font.bold = True
p.add_run(datetime.date.today().strftime('%d/%m/%Y'))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('Última actualización: ').font.bold = True
p.add_run(datetime.date.today().strftime('%d/%m/%Y'))

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('DOCUMENTO CONFIDENCIAL')
run.font.size = Pt(10)
run.font.italic = True
run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

# Salto de página
doc.add_page_break()

# ─── ÍNDICE ───────────────────────────────────────────
doc.add_heading('ÍNDICE', level=1)
indice_items = [
    '1. Identificación del Responsable del Tratamiento',
    '2. Alcance y Aceptación de la Política',
    '3. Definiciones',
    '4. Datos Personales que se Recolectan',
    '5. Finalidad del Tratamiento de Datos',
    '6. Base Legal del Tratamiento',
    '7. Almacenamiento y Seguridad de los Datos',
    '8. Compartición de Datos con Terceros',
    '9. Permisos del Dispositivo',
    '10. Tiempo de Conservación de los Datos',
    '11. Derechos del Usuario (Derechos ARCO)',
    '12. Transferencia Internacional de Datos',
    '13. Uso de Cookies (Plataforma Web)',
    '14. Datos de Menores de Edad',
    '15. Modificaciones a la Política de Privacidad',
    '16. Legislación Aplicable y Jurisdicción',
    '17. Contacto',
]
for item in indice_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ═══════════════════════════════════════════════════════
# SECCIÓN 1 — IDENTIFICACIÓN DEL RESPONSABLE
# ═══════════════════════════════════════════════════════
doc.add_heading('1. Identificación del Responsable del Tratamiento', level=1)

add_para(
    'El responsable del tratamiento de los datos personales recopilados a través de la plataforma '
    'Flow-QA/QC (en adelante, "la Plataforma") es:'
)

p = doc.add_paragraph()
p.add_run('Razón Social: ').font.bold = True
p.add_run('Vastoria Services')

p = doc.add_paragraph()
p.add_run('Representante Legal: ').font.bold = True
p.add_run('Joseph Fred Yauri Cajas')

p = doc.add_paragraph()
p.add_run('DNI: ').font.bold = True
p.add_run('72813231')

p = doc.add_paragraph()
p.add_run('Domicilio Legal: ').font.bold = True
p.add_run('Av. Ricardo Menéndez 270, Huancayo, Perú')

p = doc.add_paragraph()
p.add_run('Ciudad / País: ').font.bold = True
p.add_run('Huancayo, Perú')

p = doc.add_paragraph()
p.add_run('Correo de Privacidad: ').font.bold = True
p.add_run('vastoriaservices@gmail.com')

p = doc.add_paragraph()
p.add_run('Teléfono de contacto: ').font.bold = True
p.add_run('+51 973 785 282')

add_para(
    'En adelante, el responsable será referido como "la Empresa", "nosotros" o "Flow-QA/QC".',
    italic=True, size=10, color=RGBColor(0x64, 0x74, 0x8B), space_after=12
)

# ═══════════════════════════════════════════════════════
# SECCIÓN 2 — ALCANCE Y ACEPTACIÓN
# ═══════════════════════════════════════════════════════
doc.add_heading('2. Alcance y Aceptación de la Política', level=1)

add_para(
    'La presente Política de Privacidad (en adelante, "la Política") regula el tratamiento de los '
    'datos personales de todos los usuarios que accedan y utilicen la Plataforma Flow-QA/QC en '
    'cualquiera de sus versiones:'
)

add_bullet('Aplicación móvil para Android (React Native / Expo)', bold_prefix=None)
add_bullet('Plataforma web (Next.js)', bold_prefix=None)
add_bullet('Aplicación de escritorio (Electron)', bold_prefix=None)

add_para(
    'Al registrarse, iniciar sesión o utilizar cualquier funcionalidad de la Plataforma, el usuario '
    'declara haber leído, comprendido y aceptado íntegramente la presente Política. Si el usuario '
    'no está de acuerdo con alguna de las condiciones aquí descritas, deberá abstenerse de utilizar '
    'la Plataforma.'
)

add_para(
    'Esta Política aplica a todos los roles de usuario dentro de la Plataforma: Creator (Administrador), '
    'Resident (Residente/Jefe), Supervisor (QC) y Operator (Operador).'
)

# ═══════════════════════════════════════════════════════
# SECCIÓN 3 — DEFINICIONES
# ═══════════════════════════════════════════════════════
doc.add_heading('3. Definiciones', level=1)

definiciones = [
    ('Datos Personales', 'Toda información sobre una persona natural que la identifica o la hace identificable.'),
    ('Tratamiento', 'Cualquier operación realizada sobre datos personales: recopilación, registro, almacenamiento, uso, modificación, comunicación, eliminación.'),
    ('Titular', 'Persona natural a quien corresponden los datos personales.'),
    ('Responsable del Tratamiento', 'Persona natural o jurídica que decide sobre el tratamiento de los datos personales.'),
    ('Encargado del Tratamiento', 'Persona natural o jurídica que realiza el tratamiento de datos por cuenta del Responsable.'),
    ('Banco de Datos', 'Conjunto organizado de datos personales, automatizado o no.'),
    ('Protocolo', 'Formato digital de control de calidad que registra la conformidad de partidas ejecutadas en obra.'),
    ('Evidencia', 'Fotografía capturada in-situ como respaldo técnico de un protocolo de calidad.'),
    ('Dossier', 'Expediente documental consolidado que agrupa todos los protocolos y evidencias de un proyecto.'),
]

add_table(
    ['Término', 'Definición'],
    definiciones,
    col_widths=[2.0, 4.5]
)

# ═══════════════════════════════════════════════════════
# SECCIÓN 4 — DATOS QUE SE RECOLECTAN
# ═══════════════════════════════════════════════════════
doc.add_heading('4. Datos Personales que se Recolectan', level=1)

add_para(
    'La Plataforma recopila las siguientes categorías de datos personales, de acuerdo con la '
    'funcionalidad que el usuario utilice:'
)

doc.add_heading('4.1. Datos de Identificación y Cuenta', level=2)
add_bullet('Nombre completo (nombre y apellido)')
add_bullet('Rol asignado en el proyecto (Creator, Resident, Supervisor, Operator)')
add_bullet('Contraseña de acceso (almacenada localmente en el dispositivo)')
add_bullet('PIN de acceso rápido (opcional)')
add_bullet('Imagen de firma digital del usuario')

doc.add_heading('4.2. Datos del Dispositivo', level=2)
add_bullet('Token de notificaciones push (Expo Push Token)')
add_bullet('Plataforma del dispositivo (Android / iOS)')
add_bullet('Identificador del dispositivo (para gestión de notificaciones)')

doc.add_heading('4.3. Datos de Evidencia Técnica', level=2)
add_bullet('Fotografías capturadas exclusivamente desde la cámara del dispositivo (nunca desde galería)')
add_bullet('Metadatos de las fotografías: fecha, hora, comentario técnico, logo del proyecto')
add_bullet('Archivos técnicos subidos por el usuario: hojas de cálculo Excel, planos PDF, planos DWG')

doc.add_heading('4.4. Datos de Actividad en la Plataforma', level=2)
add_bullet('Registros de protocolos creados, llenados, enviados, aprobados y rechazados')
add_bullet('Anotaciones y observaciones realizadas sobre planos técnicos')
add_bullet('Comentarios en hilos de observaciones')
add_bullet('Historial de acciones con sellos de tiempo (timestamps)')
add_bullet('Notas de dashboard del proyecto')
add_bullet('Registros de no conformidades')

doc.add_heading('4.5. Datos de Contacto de Terceros', level=2)
add_bullet('Directorio de contactos del proyecto: nombre, teléfono y rol de los participantes')
add_para(
    'Estos datos son ingresados voluntariamente por el administrador del proyecto y corresponden '
    'a personas vinculadas al proyecto de construcción.',
    italic=True, size=10
)

# ═══════════════════════════════════════════════════════
# SECCIÓN 5 — FINALIDAD DEL TRATAMIENTO
# ═══════════════════════════════════════════════════════
doc.add_heading('5. Finalidad del Tratamiento de Datos', level=1)

add_para(
    'Los datos personales recopilados son utilizados exclusivamente para las siguientes finalidades:'
)

finalidades = [
    ('Nombre y Rol', 'Identificación del usuario y control de acceso jerárquico dentro del proyecto'),
    ('Contraseña / PIN', 'Autenticación y verificación de identidad del usuario'),
    ('Firma digital', 'Validación legal de aprobación y firma de protocolos de calidad'),
    ('Fotografías', 'Evidencia técnica de cumplimiento de estándares de calidad en obra'),
    ('Timestamps', 'Trazabilidad temporal e inmutabilidad del registro para fines de auditoría'),
    ('Push tokens', 'Envío de notificaciones operativas del proyecto (aprobaciones, rechazos, observaciones)'),
    ('Archivos técnicos', 'Gestión documental técnica del proyecto de construcción'),
    ('Datos de actividad', 'Generación de reportes, dashboards y dossier de calidad del proyecto'),
    ('Directorio de contactos', 'Facilitar la comunicación entre los participantes del proyecto'),
]

add_table(
    ['Dato', 'Finalidad'],
    finalidades,
    col_widths=[2.0, 4.5]
)

add_para(
    'Flow-QA/QC NO utiliza los datos personales para fines publicitarios, de marketing, '
    'elaboración de perfiles comerciales, ni los comparte con terceros para dichos fines.',
    bold=True, size=10
)

# ═══════════════════════════════════════════════════════
# SECCIÓN 6 — BASE LEGAL
# ═══════════════════════════════════════════════════════
doc.add_heading('6. Base Legal del Tratamiento', level=1)

add_para(
    'El tratamiento de los datos personales se fundamenta en las siguientes bases legales:'
)

doc.add_heading('6.1. Marco Normativo Nacional (Perú)', level=2)
add_bullet('Ley N° 29733', bold_prefix=None)
p = doc.add_paragraph('Ley de Protección de Datos Personales.', style='List Bullet 2')
add_bullet('Decreto Supremo N° 003-2013-JUS', bold_prefix=None)
p = doc.add_paragraph('Reglamento de la Ley de Protección de Datos Personales.', style='List Bullet 2')
add_bullet('Directiva de Seguridad de la Información', bold_prefix=None)
p = doc.add_paragraph('Aprobada por la Autoridad Nacional de Protección de Datos Personales (ANPDP).', style='List Bullet 2')

doc.add_heading('6.2. Bases Legales del Tratamiento', level=2)
add_bullet('Consentimiento explícito: ', bold_prefix=None)
p = doc.add_paragraph('Otorgado por el usuario al registrarse y aceptar la presente Política.', style='List Bullet 2')
add_bullet('Ejecución contractual: ', bold_prefix=None)
p = doc.add_paragraph('El tratamiento es necesario para la prestación del servicio contratado.', style='List Bullet 2')
add_bullet('Interés legítimo: ', bold_prefix=None)
p = doc.add_paragraph('Garantizar la seguridad, trazabilidad y calidad en los procesos constructivos.', style='List Bullet 2')
add_bullet('Obligación legal: ', bold_prefix=None)
p = doc.add_paragraph('Cumplimiento de normativas de calidad en construcción y conservación de expedientes técnicos.', style='List Bullet 2')

# ═══════════════════════════════════════════════════════
# SECCIÓN 7 — ALMACENAMIENTO Y SEGURIDAD
# ═══════════════════════════════════════════════════════
doc.add_heading('7. Almacenamiento y Seguridad de los Datos', level=1)

doc.add_heading('7.1. Ubicación del Almacenamiento', level=2)
add_para('Los datos personales se almacenan en los siguientes entornos:')

doc.add_heading('Almacenamiento Local (Dispositivo del Usuario)', level=3)
add_bullet('Base de datos SQLite (WatermelonDB) en el almacenamiento privado de la aplicación')
add_bullet('Caché local de fotografías, planos y archivos técnicos')
add_bullet('Credenciales de sesión en almacenamiento seguro del dispositivo (AsyncStorage)')

doc.add_heading('Almacenamiento en la Nube', level=3)
add_bullet('Base de datos PostgreSQL gestionada por Supabase (infraestructura AWS)')
add_bullet('Archivos (fotografías, planos, firmas) en Amazon Web Services S3 (región us-east-2, Ohio, EE.UU.)')

doc.add_heading('7.2. Medidas de Seguridad Implementadas', level=2)

medidas = [
    ('Cifrado en tránsito', 'Todas las comunicaciones utilizan protocolo HTTPS/TLS'),
    ('URLs pre-firmadas', 'Acceso a archivos mediante URLs temporales con expiración de 5 minutos'),
    ('Acceso por proyecto', 'Cada proyecto requiere contraseña para incorporación'),
    ('Control de roles', '4 niveles de permisos diferenciados (Creator, Resident, Supervisor, Operator)'),
    ('Bloqueo de galería', 'Las fotografías solo pueden capturarse desde la cámara nativa del dispositivo'),
    ('Estampado automático', 'Cada foto incluye marca de tiempo, logo y metadatos inalterables'),
    ('Sincronización segura', 'Resolución de conflictos por timestamp con verificación de integridad'),
    ('Detección de cambios', 'Sistema ETag para verificar integridad de archivos descargados'),
]

add_table(
    ['Medida', 'Descripción'],
    medidas,
    col_widths=[2.0, 4.5]
)

# ═══════════════════════════════════════════════════════
# SECCIÓN 8 — TERCEROS
# ═══════════════════════════════════════════════════════
doc.add_heading('8. Compartición de Datos con Terceros', level=1)

add_para(
    'Para la prestación del servicio, Flow-QA/QC utiliza los siguientes proveedores de servicios '
    'tecnológicos (encargados del tratamiento):'
)

terceros = [
    ('Supabase (AWS)', 'Datos del proyecto, usuarios, protocolos', 'Base de datos en la nube (PostgreSQL)'),
    ('Amazon Web Services (S3)', 'Fotografías, planos, firmas, archivos', 'Almacenamiento seguro de archivos'),
    ('Expo / Firebase (FCM)', 'Tokens de dispositivo', 'Envío de notificaciones push'),
]

add_table(
    ['Proveedor', 'Datos Compartidos', 'Finalidad'],
    terceros,
    col_widths=[1.8, 2.5, 2.2]
)

add_para('Flow-QA/QC declara expresamente que:', bold=True)
add_bullet('NO vende, alquila ni comercializa datos personales a terceros.')
add_bullet('NO comparte datos con fines publicitarios, de marketing o elaboración de perfiles.')
add_bullet('Los datos del proyecto son accesibles ÚNICAMENTE por los usuarios asignados al mismo proyecto.')
add_bullet('Los proveedores de servicios están sujetos a sus propias políticas de privacidad y cumplimiento normativo.')

# ═══════════════════════════════════════════════════════
# SECCIÓN 9 — PERMISOS DEL DISPOSITIVO
# ═══════════════════════════════════════════════════════
doc.add_heading('9. Permisos del Dispositivo', level=1)

add_para(
    'La aplicación móvil solicita los siguientes permisos del dispositivo para su correcto funcionamiento:'
)

permisos = [
    ('Cámara', 'Captura de evidencia fotográfica para protocolos y observaciones', 'Sí'),
    ('Almacenamiento', 'Caché local de planos, fotos y base de datos', 'Sí'),
    ('Internet', 'Sincronización con la nube y descarga de datos', 'No (funciona offline)'),
    ('Notificaciones', 'Alertas de aprobación, rechazo y nuevas observaciones', 'Opcional'),
]

add_table(
    ['Permiso', 'Uso', 'Obligatorio'],
    permisos,
    col_widths=[1.5, 3.5, 1.5]
)

add_para('Aclaraciones importantes:', bold=True)
add_bullet('La cámara solo captura fotografías nuevas. El sistema bloquea completamente el acceso a la galería del dispositivo para garantizar la autenticidad de la evidencia.')
add_bullet('La Plataforma NO recopila datos de geolocalización (GPS) del usuario.')
add_bullet('Los permisos pueden ser gestionados por el usuario desde la configuración de su dispositivo.')

# ═══════════════════════════════════════════════════════
# SECCIÓN 10 — TIEMPO DE CONSERVACIÓN
# ═══════════════════════════════════════════════════════
doc.add_heading('10. Tiempo de Conservación de los Datos', level=1)

add_para(
    'Los datos personales serán conservados durante el tiempo necesario para cumplir con las '
    'finalidades descritas en esta Política, de acuerdo con los siguientes criterios:'
)

p = doc.add_paragraph()
p.add_run('Datos de proyectos activos: ').font.bold = True
p.add_run('Se conservan mientras el proyecto mantenga el estado ACTIVO en la Plataforma.')

p = doc.add_paragraph()
p.add_run('Datos de proyectos cerrados: ').font.bold = True
p.add_run('Se conservarán por un período de diez (10) años contados desde la fecha de cierre del proyecto, en cumplimiento de la normativa de conservación de expedientes técnicos de construcción.')

p = doc.add_paragraph()
p.add_run('Evidencia fotográfica: ').font.bold = True
p.add_run('Se conservará por el mismo período que el proyecto, dado su carácter de respaldo técnico-legal.')

p = doc.add_paragraph()
p.add_run('Datos de cuenta de usuario: ').font.bold = True
p.add_run('Se mantienen hasta que el administrador del proyecto elimine la cuenta del usuario.')

p = doc.add_paragraph()
p.add_run('Tokens de notificaciones push: ').font.bold = True
p.add_run('Se eliminan automáticamente al cerrar sesión en el dispositivo.')

add_para(
    'Nota: En el ámbito de la construcción civil peruana, los expedientes técnicos suelen conservarse '
    'por un mínimo de diez (10) años por temas de responsabilidad civil. Se recomienda consultar con '
    'un asesor legal para definir el período exacto aplicable.',
    italic=True, size=10, color=RGBColor(0x64, 0x74, 0x8B)
)

# ═══════════════════════════════════════════════════════
# SECCIÓN 11 — DERECHOS ARCO
# ═══════════════════════════════════════════════════════
doc.add_heading('11. Derechos del Usuario (Derechos ARCO)', level=1)

add_para(
    'De conformidad con la Ley N° 29733 y su Reglamento, todo usuario titular de datos personales '
    'tiene derecho a:'
)

derechos = [
    ('Acceso', 'Solicitar información sobre qué datos personales se tienen almacenados sobre su persona, la finalidad del tratamiento y los destinatarios.'),
    ('Rectificación', 'Solicitar la corrección o actualización de datos personales que sean inexactos, incompletos o desactualizados.'),
    ('Cancelación', 'Solicitar la eliminación de sus datos personales cuando estos ya no sean necesarios para la finalidad para la que fueron recopilados.'),
    ('Oposición', 'Oponerse al tratamiento de sus datos personales por motivos legítimos y fundados.'),
]

add_table(
    ['Derecho', 'Descripción'],
    derechos,
    col_widths=[1.5, 5.0]
)

doc.add_heading('11.1. Procedimiento para Ejercer los Derechos ARCO', level=2)

add_para('Para ejercer cualquiera de estos derechos, el usuario deberá:')

add_bullet('Enviar una solicitud escrita al correo electrónico: vastoriaservices@gmail.com', bold_prefix=None)

add_bullet('Incluir en la solicitud:', bold_prefix=None)
p = doc.add_paragraph('a) Nombre completo del titular', style='List Bullet 2')
p = doc.add_paragraph('b) Descripción clara del derecho que desea ejercer', style='List Bullet 2')
p = doc.add_paragraph('c) Proyecto(s) al que está vinculado', style='List Bullet 2')
p = doc.add_paragraph('d) Documento de identidad (copia simple)', style='List Bullet 2')

doc.add_heading('11.2. Plazos de Respuesta', level=2)
add_bullet('La Empresa responderá la solicitud en un plazo máximo de diez (10) días hábiles contados desde la recepción de la solicitud completa.')
add_bullet('En caso de que la solicitud requiera información adicional, se notificará al titular dentro de los cinco (5) primeros días hábiles.')

doc.add_heading('11.3. Limitaciones', level=2)
add_para(
    'El ejercicio del derecho de cancelación u oposición podría afectar la trazabilidad e integridad '
    'del expediente técnico del proyecto. En tales casos, la Empresa informará al titular sobre las '
    'consecuencias antes de proceder, y podrá denegar la solicitud cuando exista una obligación legal '
    'de conservación de los datos.'
)

# ═══════════════════════════════════════════════════════
# SECCIÓN 12 — TRANSFERENCIA INTERNACIONAL
# ═══════════════════════════════════════════════════════
doc.add_heading('12. Transferencia Internacional de Datos', level=1)

add_para(
    'Los datos personales recopilados por Flow-QA/QC son transferidos y almacenados en servidores '
    'ubicados fuera del territorio peruano, específicamente:'
)

add_bullet('Amazon Web Services (AWS): servidores en la región us-east-2 (Ohio, Estados Unidos)')
add_bullet('Supabase: infraestructura sobre AWS con servidores en Estados Unidos')
add_bullet('Expo / Firebase Cloud Messaging: servidores de Google en Estados Unidos')

add_para(
    'Esta transferencia internacional se realiza con las siguientes garantías:'
)

add_bullet('Consentimiento explícito del usuario al aceptar la presente Política.')
add_bullet('Los proveedores cumplen con certificaciones de seguridad reconocidas internacionalmente (SOC 2, ISO 27001, entre otras).')
add_bullet('Los datos se transmiten de forma cifrada mediante protocolos HTTPS/TLS.')
add_bullet('El acceso a los datos está restringido mediante controles de autenticación y autorización.')

# ═══════════════════════════════════════════════════════
# SECCIÓN 13 — COOKIES
# ═══════════════════════════════════════════════════════
doc.add_heading('13. Uso de Cookies (Plataforma Web)', level=1)

add_para(
    'La versión web de Flow-QA/QC utiliza las siguientes cookies estrictamente necesarias para el '
    'funcionamiento de la Plataforma:'
)

cookies = [
    ('scua_user_id', 'Sesión', 'Identificación del usuario autenticado'),
    ('Cookies de TanStack Query', 'Sesión', 'Gestión de caché de datos del lado del cliente'),
]

add_table(
    ['Cookie', 'Tipo', 'Finalidad'],
    cookies,
    col_widths=[2.2, 1.5, 2.8]
)

add_para(
    'Flow-QA/QC NO utiliza cookies de seguimiento, analíticas de terceros, ni cookies publicitarias.',
    bold=True, size=10
)

# ═══════════════════════════════════════════════════════
# SECCIÓN 14 — MENORES
# ═══════════════════════════════════════════════════════
doc.add_heading('14. Datos de Menores de Edad', level=1)

add_para(
    'Flow-QA/QC es una plataforma diseñada exclusivamente para uso profesional en el ámbito de la '
    'construcción civil. No está dirigida a menores de edad y no recopila intencionalmente datos '
    'personales de personas menores de dieciocho (18) años.'
)

add_para(
    'Si la Empresa toma conocimiento de que se han recopilado datos de un menor de edad sin el '
    'consentimiento de su representante legal, procederá a eliminar dichos datos de forma inmediata.'
)

# ═══════════════════════════════════════════════════════
# SECCIÓN 15 — MODIFICACIONES
# ═══════════════════════════════════════════════════════
doc.add_heading('15. Modificaciones a la Política de Privacidad', level=1)

add_para(
    'La Empresa se reserva el derecho de modificar la presente Política de Privacidad en cualquier '
    'momento, a fin de adaptarla a cambios legislativos, jurisprudenciales, tecnológicos o de '
    'prácticas comerciales.'
)

add_para('Las modificaciones serán comunicadas a los usuarios mediante:')
add_bullet('Notificación push a través de la aplicación móvil')
add_bullet('Aviso visible en la plataforma web al iniciar sesión')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Correo electrónico (si se dispone del dato): ')
add_blank(p, 'Opcional')

add_para(
    'El uso continuado de la Plataforma después de la publicación de las modificaciones constituirá '
    'la aceptación de las mismas por parte del usuario.'
)

add_para(
    'Se mantendrá un historial de versiones de esta Política con las fechas de cada actualización.'
)

# ═══════════════════════════════════════════════════════
# SECCIÓN 16 — LEGISLACIÓN
# ═══════════════════════════════════════════════════════
doc.add_heading('16. Legislación Aplicable y Jurisdicción', level=1)

add_para(
    'La presente Política de Privacidad se rige por las leyes de la República del Perú, en particular:'
)

add_bullet('Ley N° 29733 — Ley de Protección de Datos Personales')
add_bullet('Decreto Supremo N° 003-2013-JUS — Reglamento de la Ley')
add_bullet('Constitución Política del Perú, Artículo 2, inciso 6 (derecho a la intimidad)')

p = doc.add_paragraph()
p.add_run('Para cualquier controversia derivada de la interpretación o aplicación de esta Política, '
          'las partes se someten a la jurisdicción de los jueces y tribunales de ')
p.add_run('Huancayo')
p.add_run(', República del Perú.')

# ═══════════════════════════════════════════════════════
# SECCIÓN 17 — CONTACTO
# ═══════════════════════════════════════════════════════
doc.add_heading('17. Contacto', level=1)

add_para(
    'Para cualquier consulta, solicitud o reclamo relacionado con el tratamiento de datos personales, '
    'el usuario puede comunicarse a través de los siguientes medios:'
)

p = doc.add_paragraph()
p.add_run('Correo electrónico: ').font.bold = True
p.add_run('vastoriaservices@gmail.com')

p = doc.add_paragraph()
p.add_run('Teléfono: ').font.bold = True
p.add_run('+51 973 785 282')

p = doc.add_paragraph()
p.add_run('Dirección postal: ').font.bold = True
p.add_run('Av. Ricardo Menéndez 270, Huancayo, Perú')

p = doc.add_paragraph()
p.add_run('Horario de atención: ').font.bold = True
p.add_run('Lunes a Viernes, 8:00 a.m. a 6:00 p.m.')

add_para(
    'Asimismo, el usuario tiene derecho a presentar una reclamación ante la Autoridad Nacional de '
    'Protección de Datos Personales (ANPDP) del Ministerio de Justicia y Derechos Humanos del Perú, '
    'en caso de considerar que sus derechos han sido vulnerados.'
)

# ═══════════════════════════════════════════════════════
# FIRMAS
# ═══════════════════════════════════════════════════════
doc.add_page_break()

doc.add_heading('Constancia de Aprobación', level=1)

add_para(
    'El presente documento ha sido revisado y aprobado por los abajo firmantes, entrando en vigor '
    'a partir de la fecha indicada en la portada.'
)

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

add_separator()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('Firma del Representante Legal').font.bold = True

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('_________________________________________')
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Joseph Fred Yauri Cajas')
run.font.bold = True
run.font.size = Pt(12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Representante Legal — Flow-QA/QC')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('DNI: 72813231')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('Fecha: _____ / _____ / __________').font.size = Pt(10)

doc.add_paragraph()
doc.add_paragraph()

# Nota final
add_separator()
add_para(
    'Documento oficial de Política de Privacidad de Flow-QA/QC, elaborado por Vastoria Services. '
    'Se recomienda la revisión periódica y actualización conforme a cambios normativos aplicables.',
    italic=True, size=9, color=RGBColor(0x64, 0x74, 0x8B)
)


# ─── GUARDAR ──────────────────────────────────────────
output_path = r'd:\VxP_QAQC_Automatizado\Flow-QAQC_Politica_de_Privacidad.docx'
doc.save(output_path)
print(f'Documento generado exitosamente: {output_path}')
